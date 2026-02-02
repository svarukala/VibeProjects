"""Document parsing for Word and PDF files."""

from pathlib import Path
from typing import Optional, List, Dict, Any
import re


class DocumentParser:
    """Parses Word and PDF documents to extract text content."""

    SUPPORTED_EXTENSIONS = {'.docx', '.pdf', '.doc'}

    def __init__(self):
        """Initialize the document parser."""
        self._docx_available = False
        self._pdf_available = False
        self._check_dependencies()

    def _check_dependencies(self) -> None:
        """Check if required libraries are installed."""
        try:
            import docx
            self._docx_available = True
        except ImportError:
            pass

        try:
            import pypdf
            self._pdf_available = True
        except ImportError:
            pass

    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats based on available dependencies."""
        formats = []
        if self._docx_available:
            formats.append('.docx')
        if self._pdf_available:
            formats.append('.pdf')
        return formats

    def parse(self, file_path: str) -> Dict[str, Any]:
        """Parse document and return structured content.

        Args:
            file_path: Path to the document file.

        Returns:
            dict with keys:
                - text: Full extracted text
                - sections: List of section headings and content
                - tables: Extracted tables (if any)
                - metadata: Document metadata
                - file_name: Original file name

        Raises:
            ValueError: If file type is not supported.
            FileNotFoundError: If file does not exist.
            ImportError: If required library is not installed.
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        suffix = path.suffix.lower()

        if suffix == '.docx':
            if not self._docx_available:
                raise ImportError(
                    "python-docx is required for Word documents. "
                    "Install with: pip install python-docx"
                )
            return self._parse_docx(path)
        elif suffix == '.pdf':
            if not self._pdf_available:
                raise ImportError(
                    "pypdf is required for PDF documents. "
                    "Install with: pip install pypdf"
                )
            return self._parse_pdf(path)
        elif suffix == '.doc':
            raise ValueError(
                "Legacy .doc format is not supported. "
                "Please convert to .docx format."
            )
        else:
            raise ValueError(
                f"Unsupported file type: {suffix}. "
                f"Supported formats: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )

    def _parse_docx(self, path: Path) -> Dict[str, Any]:
        """Parse Word document.

        Args:
            path: Path to the .docx file.

        Returns:
            Parsed document content.
        """
        from docx import Document

        doc = Document(str(path))

        text_parts = []
        sections = []
        tables = []
        current_section = None

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            text_parts.append(text)

            # Detect headings
            style_name = para.style.name if para.style else ''
            if style_name.startswith('Heading'):
                # Extract heading level
                level = 1
                if style_name[-1].isdigit():
                    level = int(style_name[-1])

                current_section = {
                    'level': level,
                    'title': text,
                    'content': []
                }
                sections.append(current_section)
            elif current_section is not None:
                current_section['content'].append(text)

        # Convert section content lists to strings
        for section in sections:
            section['content'] = '\n'.join(section['content'])

        # Extract tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            if table_data:
                tables.append(table_data)

        # Extract metadata
        metadata = {}
        try:
            props = doc.core_properties
            metadata = {
                'author': props.author or '',
                'title': props.title or '',
                'subject': props.subject or '',
                'created': str(props.created) if props.created else '',
                'modified': str(props.modified) if props.modified else '',
            }
        except Exception:
            pass

        return {
            'text': '\n\n'.join(text_parts),
            'sections': sections,
            'tables': tables,
            'metadata': metadata,
            'file_name': path.name,
        }

    def _parse_pdf(self, path: Path) -> Dict[str, Any]:
        """Parse PDF document.

        Args:
            path: Path to the .pdf file.

        Returns:
            Parsed document content.
        """
        from pypdf import PdfReader

        reader = PdfReader(str(path))

        text_parts = []
        for page_num, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            except Exception:
                # Skip pages that can't be extracted
                continue

        # Try to extract metadata
        metadata = {}
        try:
            if reader.metadata:
                metadata = {
                    'author': reader.metadata.author or '',
                    'title': reader.metadata.title or '',
                    'subject': reader.metadata.subject or '',
                    'creator': reader.metadata.creator or '',
                }
        except Exception:
            pass

        full_text = '\n\n'.join(text_parts)

        # Try to detect sections from text (heuristic based on common patterns)
        sections = self._detect_sections_from_text(full_text)

        return {
            'text': full_text,
            'sections': sections,
            'tables': [],  # PDF table extraction is complex, skip for now
            'metadata': metadata,
            'file_name': path.name,
            'page_count': len(reader.pages),
        }

    def _detect_sections_from_text(self, text: str) -> List[Dict[str, Any]]:
        """Detect section headings from plain text using heuristics.

        Args:
            text: Plain text content.

        Returns:
            List of detected sections.
        """
        sections = []
        lines = text.split('\n')

        # Common heading patterns
        heading_patterns = [
            r'^#+\s+(.+)$',  # Markdown style
            r'^(\d+\.?\s+[A-Z][^.]+)$',  # Numbered headings
            r'^([A-Z][A-Z\s]{2,50})$',  # ALL CAPS headings
            r'^((?:Introduction|Overview|Architecture|Summary|Conclusion|Background|Implementation|Design|Requirements|Components|Data Flow|Integration).*)$',
        ]

        current_section = None

        for line in lines:
            line = line.strip()
            if not line:
                continue

            is_heading = False
            for pattern in heading_patterns:
                match = re.match(pattern, line, re.IGNORECASE)
                if match and len(line) < 100:  # Headings are usually short
                    is_heading = True
                    current_section = {
                        'level': 1,
                        'title': line,
                        'content': []
                    }
                    sections.append(current_section)
                    break

            if not is_heading and current_section is not None:
                current_section['content'].append(line)

        # Convert content lists to strings
        for section in sections:
            section['content'] = '\n'.join(section['content'])

        return sections

    def extract_components(self, parsed_doc: Dict[str, Any]) -> Dict[str, bool]:
        """Extract M365 components from parsed document using keyword matching.

        Args:
            parsed_doc: Parsed document from parse() method.

        Returns:
            dict of component flags (e.g., {'has_copilot': True, ...})
        """
        text = parsed_doc['text'].lower()

        # Also check tables for component mentions
        table_text = ''
        for table in parsed_doc.get('tables', []):
            for row in table:
                table_text += ' '.join(row).lower() + ' '

        combined_text = text + ' ' + table_text

        return {
            'has_user': any(w in combined_text for w in [
                'user', 'employee', 'end user', 'worker', 'customer', 'client'
            ]),
            'has_copilot': (
                'copilot' in combined_text and
                'copilot studio' not in combined_text
            ) or 'm365 copilot' in combined_text or 'microsoft 365 copilot' in combined_text,
            'has_copilot_studio': any(w in combined_text for w in [
                'copilot studio', 'custom agent', 'copilot agent'
            ]),
            'has_sharepoint': any(w in combined_text for w in [
                'sharepoint', 'share point', 'document library', 'sp online'
            ]),
            'has_teams': any(w in combined_text for w in [
                'microsoft teams', 'ms teams', 'teams channel', 'teams chat'
            ]) or (combined_text.count('teams') > 1 and 'microsoft' in combined_text),
            'has_graph_api': any(w in combined_text for w in [
                'graph api', 'microsoft graph', 'msgraph', 'ms graph'
            ]),
            'has_dataverse': any(w in combined_text for w in [
                'dataverse', 'common data service', 'cds'
            ]),
            'has_power_automate': any(w in combined_text for w in [
                'power automate', 'powerautomate', 'cloud flow', 'automated flow',
                'workflow automation'
            ]),
            'has_power_apps': any(w in combined_text for w in [
                'power apps', 'powerapps', 'canvas app', 'model-driven app',
                'model driven app'
            ]),
            'has_power_bi': any(w in combined_text for w in [
                'power bi', 'powerbi', 'bi dashboard', 'bi report'
            ]),
            'has_azure_openai': any(w in combined_text for w in [
                'azure openai', 'gpt-4', 'gpt-3.5', 'gpt4', 'gpt35',
                'large language model', 'llm', 'openai service'
            ]),
            'has_azure_ai_search': any(w in combined_text for w in [
                'ai search', 'cognitive search', 'azure search',
                'vector search', 'semantic search', 'search index'
            ]),
            'has_azure_foundry': any(w in combined_text for w in [
                'ai foundry', 'azure ai foundry', 'prompt flow',
                'model catalog', 'azure machine learning'
            ]),
            'has_connector': any(w in combined_text for w in [
                'connector', 'custom connector', 'api connection',
                'third-party', 'third party', 'webhook', 'rest api'
            ]),
            'has_sql': any(w in combined_text for w in [
                'sql server', 'azure sql', 'sql database', 'mssql'
            ]),
            'has_blob_storage': any(w in combined_text for w in [
                'blob storage', 'azure storage', 'storage account',
                'file storage', 'azure blob'
            ]),
            'has_logic_apps': any(w in combined_text for w in [
                'logic apps', 'logic app', 'azure logic'
            ]),
            'has_functions': any(w in combined_text for w in [
                'azure function', 'azure functions', 'serverless function'
            ]),
        }

    def get_summary(self, parsed_doc: Dict[str, Any], max_length: int = 500) -> str:
        """Get a summary of the document for display.

        Args:
            parsed_doc: Parsed document.
            max_length: Maximum length of summary.

        Returns:
            Summary string.
        """
        text = parsed_doc['text']
        if len(text) <= max_length:
            return text

        # Try to end at a sentence boundary
        truncated = text[:max_length]
        last_period = truncated.rfind('.')
        if last_period > max_length // 2:
            truncated = truncated[:last_period + 1]

        return truncated + '...'
