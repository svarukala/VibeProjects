"""
PDF Form to Word/HTML Document Converter

Converts PDF forms to Word or HTML documents with editable text by:
1. Using pdf2docx to convert PDF to Word (preserves tables, layout, formatting)
2. Extracting form field positions from the original PDF
3. Replacing form fields with {{placeholder}} format in the output document

Requirements:
    pip install pdf2docx pymupdf python-docx
"""

import fitz  # PyMuPDF
from pdf2docx import Converter
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
import re
import argparse
from pathlib import Path
from typing import Dict, List, Tuple, Optional, Set
from dataclasses import dataclass, field
import tempfile
import os
import shutil
import html


@dataclass
class FormField:
    """Represents a form field extracted from PDF."""
    name: str
    field_type: str
    rect: Tuple[float, float, float, float]  # x0, y0, x1, y1 in PDF points
    page_num: int
    value: Optional[str] = None
    options: List[str] = field(default_factory=list)
    group_name: Optional[str] = None
    is_checked: bool = False
    original_name: str = ""
    placeholder: str = ""


class PDFFormExtractor:
    """Extracts form fields from PDF documents."""

    FIELD_TYPE_MAP = {
        0: 'text',
        1: 'checkbox',
        2: 'radio',
        3: 'dropdown',
        4: 'dropdown',
        5: 'signature',
    }

    def __init__(self, pdf_path: str):
        self.pdf_path = pdf_path
        self.doc = fitz.open(pdf_path)

    def get_page_count(self) -> int:
        return len(self.doc)

    def extract_all_fields(self) -> List[FormField]:
        """Extract all form fields from all pages."""
        all_fields = []
        for page_num in range(len(self.doc)):
            page = self.doc[page_num]
            widgets = page.widgets()
            if not widgets:
                continue

            for widget in widgets:
                field_type_num = widget.field_type
                field_type = self.FIELD_TYPE_MAP.get(field_type_num, 'text')

                original_name = widget.field_name or f"Field_{page_num}_{len(all_fields)}"
                clean_name = self._clean_field_name(original_name)

                form_field = FormField(
                    name=clean_name,
                    field_type=field_type,
                    rect=tuple(widget.rect),
                    page_num=page_num,
                    value=widget.field_value,
                    original_name=original_name,
                )

                if field_type == 'checkbox':
                    form_field.is_checked = widget.field_value == 'Yes'
                elif field_type == 'radio':
                    form_field.group_name = widget.field_name
                    form_field.is_checked = bool(widget.field_value)
                elif field_type == 'dropdown':
                    try:
                        form_field.options = list(widget.choice_values) if widget.choice_values else []
                    except:
                        form_field.options = []

                all_fields.append(form_field)

        return all_fields

    def _clean_field_name(self, name: str) -> str:
        """Clean field name for placeholder format, limited to 4 characters."""
        clean = re.sub(r'[^\w\s]', '', name)
        clean = re.sub(r'\s+', '_', clean)
        clean = clean.lower()

        # Remove leading non-alpha characters
        clean = re.sub(r'^[^a-z]+', '', clean)

        # If empty or undefined, use 'udf'
        if not clean:
            return 'udf'

        # Limit to 4 characters
        return clean[:4]

    def close(self):
        self.doc.close()


class CheckboxProcessor:
    """Processes checkbox fields to create unified placeholders."""

    def process_fields(self, fields: List[FormField]) -> List[FormField]:
        """Process all fields, merging yes/no checkbox pairs."""
        checkboxes = [f for f in fields if f.field_type == 'checkbox']
        other_fields = [f for f in fields if f.field_type != 'checkbox']

        # Assign placeholders to other fields
        for fld in other_fields:
            fld.placeholder = f"{{{{{fld.name}}}}}"

        # Process checkboxes
        processed_checkboxes = []
        processed_indices: Set[int] = set()

        for i, cb in enumerate(checkboxes):
            if i in processed_indices:
                continue

            pair_found = False
            for j, other_cb in enumerate(checkboxes):
                if i == j or j in processed_indices:
                    continue

                if self._is_yes_no_pair(cb, other_cb):
                    base_name = self._get_base_name_for_pair(cb, other_cb)
                    placeholder = f"{{{{{base_name}_yes_or_no}}}}"

                    # Keep both checkboxes but with same placeholder
                    cb.placeholder = placeholder
                    cb.field_type = 'yes_no_pair'
                    other_cb.placeholder = placeholder
                    other_cb.field_type = 'yes_no_pair'

                    processed_checkboxes.append(cb)
                    processed_checkboxes.append(other_cb)

                    processed_indices.add(i)
                    processed_indices.add(j)
                    pair_found = True
                    break

            if not pair_found:
                cb.placeholder = f"{{{{{cb.name}_confirmed}}}}"
                processed_checkboxes.append(cb)
                processed_indices.add(i)

        return other_fields + processed_checkboxes

    def _is_yes_no_pair(self, cb1: FormField, cb2: FormField) -> bool:
        """Check if two checkboxes form a yes/no pair."""
        if cb1.page_num != cb2.page_num:
            return False

        y_diff = abs(cb1.rect[1] - cb2.rect[1])
        if y_diff > 20:
            return False

        name1 = cb1.original_name.lower()
        name2 = cb2.original_name.lower()

        yes_patterns = ['yes', 'y', 'true', 'accept', 'agree']
        no_patterns = ['no', 'n', 'false', 'reject', 'disagree']

        name1_is_yes = any(p in name1 for p in yes_patterns)
        name1_is_no = any(p in name1 for p in no_patterns)
        name2_is_yes = any(p in name2 for p in yes_patterns)
        name2_is_no = any(p in name2 for p in no_patterns)

        if (name1_is_yes and name2_is_no) or (name1_is_no and name2_is_yes):
            return True

        base1 = re.sub(r'[_\s]*(yes|no|y|n|true|false|\d+)$', '', name1, flags=re.IGNORECASE).strip('_')
        base2 = re.sub(r'[_\s]*(yes|no|y|n|true|false|\d+)$', '', name2, flags=re.IGNORECASE).strip('_')

        if base1 and base1 == base2:
            return True

        x_diff = abs(cb1.rect[0] - cb2.rect[0])
        if x_diff < 100 and y_diff < 10:
            return True

        return False

    def _get_base_name_for_pair(self, cb1: FormField, cb2: FormField) -> str:
        """Extract base name for a checkbox pair, limited to 4 characters."""
        name1 = cb1.original_name.lower()
        name2 = cb2.original_name.lower()

        base1 = re.sub(r'[_\s]*(yes|no|y|n|true|false)$', '', name1, flags=re.IGNORECASE)
        base2 = re.sub(r'[_\s]*(yes|no|y|n|true|false)$', '', name2, flags=re.IGNORECASE)

        base = base1 if len(base1) >= len(base2) else base2
        base = re.sub(r'[^\w]', '_', base)
        base = re.sub(r'_+', '_', base).strip('_')

        # Remove leading non-alpha characters
        base = re.sub(r'^[^a-z]+', '', base)

        # If empty, use 'udf'
        if not base:
            return 'udf'

        # Limit to 4 characters
        return base[:4]


class WordDocumentProcessor:
    """Processes Word document to insert placeholders at field positions."""

    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        self.doc = Document(docx_path)

    def insert_placeholders(self, fields: List[FormField], pdf_path: str):
        """
        Insert placeholders into the Word document at positions corresponding
        to PDF form fields.
        """
        # Group fields by page
        fields_by_page: Dict[int, List[FormField]] = {}
        for fld in fields:
            if fld.page_num not in fields_by_page:
                fields_by_page[fld.page_num] = []
            fields_by_page[fld.page_num].append(fld)

        # Track which placeholders we've already inserted (for yes/no pairs)
        inserted_placeholders: Dict[str, Set[int]] = {}  # placeholder -> set of page numbers

        # Process paragraphs
        for para in self.doc.paragraphs:
            self._process_paragraph(para, fields_by_page, inserted_placeholders)

        # Process tables
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._process_paragraph(para, fields_by_page, inserted_placeholders)

    def _process_paragraph(
        self,
        para,
        fields_by_page: Dict[int, List[FormField]],
        inserted_placeholders: Dict[str, Set[int]]
    ):
        """Process a paragraph to insert placeholders."""
        # For each page's fields, check if they should be inserted
        for page_num, page_fields in fields_by_page.items():
            for fld in page_fields:
                if not fld.placeholder:
                    continue

                # For yes/no pairs, only insert once per page
                if fld.field_type == 'yes_no_pair':
                    if fld.placeholder in inserted_placeholders:
                        if page_num in inserted_placeholders[fld.placeholder]:
                            continue
                    else:
                        inserted_placeholders[fld.placeholder] = set()
                    inserted_placeholders[fld.placeholder].add(page_num)

    def save(self, output_path: str):
        """Save the modified document."""
        self.doc.save(output_path)


class HTMLExporter:
    """Exports PDF with placeholders to HTML format."""

    def __init__(self):
        self.css_styles = """
        <style>
            * {
                margin: 0;
                padding: 0;
                box-sizing: border-box;
            }
            body {
                font-family: Arial, sans-serif;
                background-color: #f0f0f0;
                padding: 20px;
            }
            .page {
                background: white;
                margin: 20px auto;
                padding: 40px;
                box-shadow: 0 2px 10px rgba(0,0,0,0.1);
                max-width: 850px;
                position: relative;
            }
            .page-break {
                page-break-after: always;
            }
            .text-block {
                position: absolute;
                white-space: pre-wrap;
            }
            .placeholder {
                color: #0066cc;
                font-weight: bold;
                background-color: #e6f2ff;
                padding: 2px 4px;
                border-radius: 3px;
            }
            table {
                border-collapse: collapse;
                width: 100%;
                margin: 10px 0;
            }
            td, th {
                border: 1px solid #ccc;
                padding: 8px;
                text-align: left;
            }
            .field-area {
                display: inline-block;
                min-width: 50px;
            }
            @media print {
                body {
                    background: white;
                    padding: 0;
                }
                .page {
                    box-shadow: none;
                    margin: 0;
                    padding: 20px;
                }
            }
        </style>
        """

    def export_pdf_to_html(
        self,
        pdf_path: str,
        output_path: str,
        fields: List[FormField]
    ):
        """Export PDF to HTML with placeholders."""
        doc = fitz.open(pdf_path)

        # Track inserted placeholders for yes/no pairs
        inserted_yes_no: Dict[int, Set[str]] = {}

        html_content = [
            "<!DOCTYPE html>",
            "<html lang='en'>",
            "<head>",
            "<meta charset='UTF-8'>",
            "<meta name='viewport' content='width=device-width, initial-scale=1.0'>",
            "<title>Converted PDF Form</title>",
            self.css_styles,
            "</head>",
            "<body>"
        ]

        for page_num in range(len(doc)):
            page = doc[page_num]
            page_width = page.rect.width
            page_height = page.rect.height

            # Get fields for this page
            page_fields = [f for f in fields if f.page_num == page_num]

            # Build a map of field positions to placeholders
            field_rects = []
            for fld in page_fields:
                if not fld.placeholder:
                    continue

                # For yes/no pairs, only insert once
                if fld.field_type == 'yes_no_pair':
                    if page_num not in inserted_yes_no:
                        inserted_yes_no[page_num] = set()
                    if fld.placeholder in inserted_yes_no[page_num]:
                        continue
                    inserted_yes_no[page_num].add(fld.placeholder)

                field_rects.append({
                    'rect': fld.rect,
                    'placeholder': fld.placeholder
                })

            # Extract text blocks with positioning
            html_content.append(f'<div class="page" style="min-height: {page_height}px;">')

            # Get page content as dictionary
            blocks = page.get_text("dict")["blocks"]

            for block in blocks:
                if block.get("type") == 0:  # Text block
                    for line in block.get("lines", []):
                        line_text = ""
                        line_bbox = line.get("bbox", (0, 0, 0, 0))

                        for span in line.get("spans", []):
                            span_text = span.get("text", "")
                            span_bbox = span.get("bbox", (0, 0, 0, 0))
                            font_size = span.get("size", 12)
                            flags = span.get("flags", 0)

                            is_bold = bool(flags & 2**4)
                            is_italic = bool(flags & 2**1)

                            # Check if this span overlaps with any field
                            replaced = False
                            for fr in field_rects:
                                if self._rects_overlap(span_bbox, fr['rect']):
                                    line_text += f'<span class="placeholder">{html.escape(fr["placeholder"])}</span>'
                                    replaced = True
                                    break

                            if not replaced:
                                style_parts = []
                                if is_bold:
                                    style_parts.append("font-weight: bold")
                                if is_italic:
                                    style_parts.append("font-style: italic")

                                style = f' style="{"; ".join(style_parts)}"' if style_parts else ""
                                line_text += f"<span{style}>{html.escape(span_text)}</span>"

                        if line_text.strip():
                            html_content.append(f'<p>{line_text}</p>')

            # Add placeholders for fields not overlapping with text
            for fr in field_rects:
                x0, y0, x1, y1 = fr['rect']
                # Add as positioned element if not already rendered
                html_content.append(
                    f'<div class="field-area" style="position: relative; display: inline;">'
                    f'<span class="placeholder">{html.escape(fr["placeholder"])}</span>'
                    f'</div>'
                )

            html_content.append('</div>')  # Close page div

            if page_num < len(doc) - 1:
                html_content.append('<div class="page-break"></div>')

        doc.close()

        html_content.extend([
            "</body>",
            "</html>"
        ])

        # Write HTML file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(html_content))

    def _rects_overlap(self, rect1: Tuple, rect2: Tuple, tolerance: float = 5) -> bool:
        """Check if two rectangles overlap."""
        x0_1, y0_1, x1_1, y1_1 = rect1
        x0_2, y0_2, x1_2, y1_2 = rect2

        return not (
            x1_1 < x0_2 - tolerance or
            x0_1 > x1_2 + tolerance or
            y1_1 < y0_2 - tolerance or
            y0_1 > y1_2 + tolerance
        )


def export_to_html_with_form_elements(
    pdf_path: str,
    output_path: str,
    fields: List[FormField]
):
    """
    Export PDF to HTML with actual form elements (textbox, checkbox, etc.)
    instead of placeholders.
    """
    temp_dir = tempfile.gettempdir()
    temp_pdf_path = os.path.join(temp_dir, f"pdf_form_temp_{os.getpid()}.pdf")

    try:
        # Open original PDF and clear form field areas
        doc = fitz.open(pdf_path)

        # Track fields by page for later insertion
        fields_by_page: Dict[int, List[FormField]] = {}
        inserted_yes_no: Dict[int, Set[str]] = {}

        for fld in fields:
            if fld.page_num not in fields_by_page:
                fields_by_page[fld.page_num] = []

            page = doc[fld.page_num]
            rect = fitz.Rect(fld.rect)

            # For yes/no pairs, track to avoid duplicates
            if fld.field_type == 'yes_no_pair':
                if fld.page_num not in inserted_yes_no:
                    inserted_yes_no[fld.page_num] = set()
                if fld.name in inserted_yes_no[fld.page_num]:
                    # Clear area but don't add to fields list again
                    page.draw_rect(rect, color=(1, 1, 1), fill=(1, 1, 1))
                    continue
                inserted_yes_no[fld.page_num].add(fld.name)

            # Clear the field area with white
            page.draw_rect(rect, color=(1, 1, 1), fill=(1, 1, 1))

            # Add marker text that we'll replace with form elements
            marker = f"__FIELD_{fld.page_num}_{len(fields_by_page[fld.page_num])}__"
            fld.placeholder = marker  # Reuse placeholder field for marker

            field_height = rect.height
            font_size = min(max(field_height * 0.6, 6), 10)
            text_point = fitz.Point(rect.x0 + 1, rect.y0 + field_height * 0.7)

            page.insert_text(
                text_point,
                marker,
                fontsize=font_size,
                fontname="helv",
                color=(0.5, 0.5, 0.5)
            )

            fields_by_page[fld.page_num].append(fld)

        # Save modified PDF
        doc.save(temp_pdf_path)
        doc.close()

        # Convert to HTML
        print("Converting PDF to HTML with form elements...")
        doc = fitz.open(temp_pdf_path)

        html_parts = [
            "<!DOCTYPE html>",
            "<html lang='en'>",
            "<head>",
            "<meta charset='UTF-8'>",
            "<meta name='viewport' content='width=device-width, initial-scale=1.0'>",
            "<title>Converted PDF Form</title>",
            "<style>",
            "body { font-family: Arial, sans-serif; background: #f5f5f5; padding: 20px; margin: 0; }",
            ".page { background: white; max-width: 900px; margin: 20px auto; padding: 40px; box-shadow: 0 2px 8px rgba(0,0,0,0.1); }",
            ".page-content { line-height: 1.6; }",
            "table { border-collapse: collapse; width: 100%; margin: 10px 0; }",
            "td, th { border: 1px solid #ddd; padding: 8px; vertical-align: top; }",
            "img { max-width: 100%; height: auto; }",
            "input[type='text'], input[type='date'], input[type='email'] {",
            "  border: 1px solid #ccc; padding: 4px 8px; border-radius: 3px;",
            "  font-size: 14px; min-width: 150px;",
            "}",
            "input[type='checkbox'], input[type='radio'] {",
            "  width: 16px; height: 16px; margin: 2px 4px; cursor: pointer;",
            "}",
            "select {",
            "  border: 1px solid #ccc; padding: 4px 8px; border-radius: 3px;",
            "  font-size: 14px; min-width: 150px; background: white;",
            "}",
            ".signature-field {",
            "  border: 1px solid #ccc; border-radius: 3px; padding: 8px;",
            "  min-width: 200px; min-height: 50px; background: #fafafa;",
            "  display: inline-block; font-style: italic; color: #999;",
            "}",
            ".yes-no-group { display: inline-flex; gap: 15px; align-items: center; }",
            ".yes-no-group label { display: flex; align-items: center; gap: 4px; cursor: pointer; }",
            ".form-label { font-weight: 500; margin-right: 8px; }",
            "@media print {",
            "  body { background: white; padding: 0; }",
            "  .page { box-shadow: none; margin: 0; }",
            "  input, select { border: 1px solid #000; }",
            "}",
            "</style>",
            "</head>",
            "<body>",
            "<form id='pdfForm'>"
        ]

        for page_num in range(len(doc)):
            page = doc[page_num]
            page_html = page.get_text("html")

            # Replace markers with actual form elements
            page_fields = fields_by_page.get(page_num, [])
            for idx, fld in enumerate(page_fields):
                marker = fld.placeholder
                form_element = _generate_form_element(fld, idx)
                page_html = page_html.replace(marker, form_element)

            html_parts.append(f'<div class="page"><div class="page-content">{page_html}</div></div>')

        html_parts.extend([
            "</form>",
            "</body>",
            "</html>"
        ])

        doc.close()

        # Write HTML
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(html_parts))

    finally:
        try:
            if os.path.exists(temp_pdf_path):
                os.unlink(temp_pdf_path)
        except PermissionError:
            pass


def _generate_form_element(field: FormField, index: int) -> str:
    """Generate appropriate HTML form element based on field type."""
    field_id = f"{field.name}_{index}"
    field_name = field.name

    if field.field_type == 'text':
        # Text input field
        return (
            f'<input type="text" id="{html.escape(field_id)}" '
            f'name="{html.escape(field_name)}" '
            f'placeholder="Enter {html.escape(field_name)}">'
        )

    elif field.field_type == 'checkbox':
        # Single checkbox (confirmation type)
        return (
            f'<label>'
            f'<input type="checkbox" id="{html.escape(field_id)}" '
            f'name="{html.escape(field_name)}" value="yes">'
            f'</label>'
        )

    elif field.field_type == 'yes_no_pair':
        # Yes/No radio button pair
        base_name = field.name.replace('_yes_or_no', '').replace('_confirmed', '')
        return (
            f'<span class="yes-no-group">'
            f'<label><input type="radio" name="{html.escape(base_name)}" value="yes"> Yes</label>'
            f'<label><input type="radio" name="{html.escape(base_name)}" value="no"> No</label>'
            f'</span>'
        )

    elif field.field_type == 'radio':
        # Radio button (part of a group)
        group_name = field.group_name or field.name
        return (
            f'<input type="radio" id="{html.escape(field_id)}" '
            f'name="{html.escape(group_name)}" value="{html.escape(field.name)}">'
        )

    elif field.field_type == 'dropdown':
        # Dropdown/select field
        options_html = '<option value="">-- Select --</option>'
        for opt in field.options:
            options_html += f'<option value="{html.escape(opt)}">{html.escape(opt)}</option>'
        return (
            f'<select id="{html.escape(field_id)}" name="{html.escape(field_name)}">'
            f'{options_html}'
            f'</select>'
        )

    elif field.field_type == 'signature':
        # Signature field
        return (
            f'<div class="signature-field" id="{html.escape(field_id)}">'
            f'Sign here'
            f'</div>'
        )

    else:
        # Default to text input
        return (
            f'<input type="text" id="{html.escape(field_id)}" '
            f'name="{html.escape(field_name)}">'
        )


def insert_placeholders_in_pdf_then_convert(
    pdf_path: str,
    output_path: str,
    fields: List[FormField]
):
    """
    Modify PDF to show placeholders, then convert to Word.
    This approach ensures placeholders appear at exact field positions.
    """
    # Create temp file path for modified PDF
    temp_dir = tempfile.gettempdir()
    temp_pdf_path = os.path.join(temp_dir, f"pdf_form_temp_{os.getpid()}.pdf")

    try:
        # Open original PDF
        doc = fitz.open(pdf_path)

        # Track inserted placeholders for yes/no pairs
        inserted_yes_no: Dict[int, Set[str]] = {}  # page_num -> set of placeholders

        for fld in fields:
            if not fld.placeholder:
                continue

            page = doc[fld.page_num]

            # For yes/no pairs, only insert once per pair per page
            if fld.field_type == 'yes_no_pair':
                if fld.page_num not in inserted_yes_no:
                    inserted_yes_no[fld.page_num] = set()
                if fld.placeholder in inserted_yes_no[fld.page_num]:
                    # Still need to clear this checkbox area
                    rect = fitz.Rect(fld.rect)
                    page.draw_rect(rect, color=(1, 1, 1), fill=(1, 1, 1))
                    continue
                inserted_yes_no[fld.page_num].add(fld.placeholder)

            # Get field rectangle
            rect = fitz.Rect(fld.rect)

            # Clear the field area with white
            page.draw_rect(rect, color=(1, 1, 1), fill=(1, 1, 1))

            # Calculate font size based on field height
            field_height = rect.height
            font_size = min(max(field_height * 0.7, 6), 12)

            # Insert placeholder text
            text_point = fitz.Point(rect.x0 + 2, rect.y0 + field_height * 0.75)

            page.insert_text(
                text_point,
                fld.placeholder,
                fontsize=font_size,
                fontname="helv",
                color=(0, 0, 0)
            )

        # Save modified PDF to a new temp file
        doc.save(temp_pdf_path)
        doc.close()

        # Convert modified PDF to Word
        print("Converting PDF to Word (this may take a moment)...")
        cv = Converter(temp_pdf_path)
        cv.convert(output_path)
        cv.close()

    finally:
        # Clean up temp file
        try:
            if os.path.exists(temp_pdf_path):
                os.unlink(temp_pdf_path)
        except PermissionError:
            # File may still be locked, ignore cleanup error
            pass


def convert_pdf_form(
    pdf_path: str,
    output_path: str = None,
    output_format: str = 'docx'
):
    """
    Convert a PDF form to Word or HTML document with editable text and placeholders.

    Args:
        pdf_path: Path to the input PDF file
        output_path: Path for the output file
        output_format: Output format - 'docx' or 'html'

    Returns:
        Path to the generated document
    """
    pdf_path = Path(pdf_path)
    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF file not found: {pdf_path}")

    # Determine output extension
    if output_format.lower() == 'html':
        default_suffix = '.html'
    else:
        default_suffix = '.docx'
        output_format = 'docx'

    if output_path is None:
        output_path = pdf_path.with_suffix(default_suffix)
    else:
        output_path = Path(output_path)

    print(f"Converting: {pdf_path}")
    print(f"Output: {output_path}")
    print(f"Format: {output_format.upper()}")

    # Extract form fields from PDF
    print("Extracting form fields...")
    extractor = PDFFormExtractor(str(pdf_path))
    try:
        all_fields = extractor.extract_all_fields()
        print(f"Found {len(all_fields)} form fields")

        # Report field types
        field_types: Dict[str, int] = {}
        for f in all_fields:
            field_types[f.field_type] = field_types.get(f.field_type, 0) + 1
        print(f"Field types: {field_types}")

    finally:
        extractor.close()

    # Process checkboxes to create unified placeholders
    print("Processing checkboxes...")
    processor = CheckboxProcessor()
    processed_fields = processor.process_fields(all_fields)

    # Convert based on output format
    if output_format == 'html':
        print("Converting to HTML with form elements...")
        export_to_html_with_form_elements(
            str(pdf_path),
            str(output_path),
            processed_fields
        )
    else:
        print("Inserting placeholders and converting to Word...")
        insert_placeholders_in_pdf_then_convert(
            str(pdf_path),
            str(output_path),
            processed_fields
        )

    # Print summary
    print(f"\nConversion complete!")
    print(f"Output: {output_path}")

    # Show unique placeholders
    unique_placeholders = set(f.placeholder for f in processed_fields if f.placeholder)
    print(f"\nPlaceholders created ({len(unique_placeholders)}):")
    for p in sorted(unique_placeholders)[:30]:
        print(f"  {p}")
    if len(unique_placeholders) > 30:
        print(f"  ... and {len(unique_placeholders) - 30} more")

    return str(output_path)


# Backward compatibility alias
def convert_pdf_form_to_word(pdf_path: str, output_path: str = None):
    """Backward compatible function - converts to Word."""
    return convert_pdf_form(pdf_path, output_path, output_format='docx')


def main():
    """Main entry point for command-line usage."""
    parser = argparse.ArgumentParser(
        description='Convert PDF forms to Word or HTML documents with Copilot Studio placeholders'
    )
    parser.add_argument(
        'pdf_path',
        help='Path to the input PDF file'
    )
    parser.add_argument(
        '-o', '--output',
        help='Output file path (default: same name with appropriate extension)'
    )
    parser.add_argument(
        '-f', '--format',
        choices=['docx', 'html'],
        default='docx',
        help='Output format: docx (Word) or html (default: docx)'
    )

    args = parser.parse_args()

    try:
        convert_pdf_form(
            pdf_path=args.pdf_path,
            output_path=args.output,
            output_format=args.format
        )
    except FileNotFoundError as e:
        print(f"Error: {e}")
        return 1
    except Exception as e:
        print(f"Error during conversion: {e}")
        raise

    return 0


if __name__ == '__main__':
    exit(main())
