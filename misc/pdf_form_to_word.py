"""
PDF Form to Word Document Converter

Converts a PDF document containing form fields (text inputs, checkboxes, 
dropdowns, radio buttons) into a Word document with Copilot Studio 
placeholder format: {{FieldName}}

Usage:
    python pdf_form_to_word.py input.pdf output.docx
"""

import sys
import re
from pathlib import Path

try:
    import pymupdf as fitz  # PyMuPDF
except ImportError:
    import fitz

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def sanitize_text(text: str) -> str:
    """
    Remove or replace control characters that are not XML-compatible.
    """
    if not text:
        return ""
    # Remove NULL bytes and control characters (except tab, newline, carriage return)
    cleaned = ''.join(char for char in text if ord(char) >= 32 or char in '\t\n\r')
    return cleaned


def sanitize_field_name(name: str) -> str:
    """
    Sanitize field name for Copilot Studio placeholder format.
    - Remove spaces and special characters
    - Convert to PascalCase
    """
    if not name:
        return "UnnamedField"
    
    # Remove special characters except alphanumeric and spaces
    cleaned = re.sub(r'[^a-zA-Z0-9\s_]', '', name)
    
    # Split by spaces or underscores and convert to PascalCase
    words = re.split(r'[\s_]+', cleaned)
    pascal_case = ''.join(word.capitalize() for word in words if word)
    
    return pascal_case if pascal_case else "UnnamedField"


def format_placeholder(field_name: str, field_type: str = None) -> str:
    """
    Format field name as Copilot Studio placeholder: {{FieldName}}
    """
    sanitized = sanitize_field_name(field_name)
    return f"{{{{{sanitized}}}}}"


def get_field_type_description(field_type: int) -> str:
    """Map PDF field type to human-readable description."""
    field_types = {
        0: "unknown",
        1: "button",  # includes checkboxes and radio buttons
        2: "text",
        3: "choice",  # dropdowns and listboxes
        4: "signature",
    }
    return field_types.get(field_type, "unknown")


def extract_form_fields(pdf_path: str) -> list:
    """
    Extract all form fields from a PDF document.
    Returns a list of field dictionaries with metadata.
    """
    doc = fitz.open(pdf_path)
    fields = []
    field_counter = {}
    
    for page_num in range(len(doc)):
        page = doc[page_num]
        
        # Get all widgets (form fields) on the page
        for widget in page.widgets():
            field_name = widget.field_name or ""
            field_type = widget.field_type
            field_type_desc = get_field_type_description(field_type)
            
            # Handle field naming for duplicates
            base_name = field_name if field_name else f"Field_{page_num + 1}"
            if base_name in field_counter:
                field_counter[base_name] += 1
            else:
                field_counter[base_name] = 1
            
            # Get field value and options
            field_value = widget.field_value
            field_flags = widget.field_flags
            
            # Determine specific field subtype
            subtype = "text"
            options = []
            
            if field_type == 1:  # Button type
                # Check if it's a checkbox or radio button
                if widget.field_flags & 32768:  # Radio button flag
                    subtype = "radio"
                elif widget.field_flags & 16384:  # Pushbutton flag
                    subtype = "button"
                else:
                    subtype = "checkbox"
            elif field_type == 2:  # Text
                subtype = "text"
            elif field_type == 3:  # Choice
                # Check if combo box (dropdown) or listbox
                if widget.field_flags & 131072:  # Combo flag
                    subtype = "dropdown"
                else:
                    subtype = "listbox"
                # Try to get choice options
                try:
                    options = widget.choice_values or []
                except:
                    options = []
            
            # Get field rectangle (position)
            rect = widget.rect
            
            field_info = {
                "name": field_name,
                "sanitized_name": sanitize_field_name(field_name),
                "page": page_num + 1,
                "type": field_type_desc,
                "subtype": subtype,
                "value": field_value,
                "options": options,
                "rect": {
                    "x0": rect.x0,
                    "y0": rect.y0,
                    "x1": rect.x1,
                    "y1": rect.y1
                },
                "placeholder": format_placeholder(field_name)
            }
            fields.append(field_info)
    
    doc.close()
    return fields


def extract_pdf_text_with_fields(pdf_path: str, fields: list) -> list:
    """
    Extract text from PDF and identify where form fields should be placed.
    Returns a list of content blocks (text and field placeholders) per page.
    """
    doc = fitz.open(pdf_path)
    pages_content = []
    
    for page_num in range(len(doc)):
        page = doc[page_num]
        page_fields = [f for f in fields if f["page"] == page_num + 1]
        
        # Get text blocks with positions
        text_dict = page.get_text("dict")
        blocks = text_dict.get("blocks", [])
        
        content_items = []
        
        for block in blocks:
            if "lines" in block:  # Text block
                for line in block["lines"]:
                    line_text = ""
                    for span in line["spans"]:
                        line_text += span["text"]
                    if line_text.strip():
                        content_items.append({
                            "type": "text",
                            "content": sanitize_text(line_text),
                            "y": line["bbox"][1]
                        })
        
        # Add field placeholders with their positions
        for field in page_fields:
            content_items.append({
                "type": "field",
                "field": field,
                "y": field["rect"]["y0"]
            })
        
        # Sort by vertical position
        content_items.sort(key=lambda x: x["y"])
        pages_content.append(content_items)
    
    doc.close()
    return pages_content


def create_word_document(pdf_path: str, output_path: str, include_field_list: bool = True):
    """
    Create a Word document from PDF form with Copilot Studio placeholders.
    """
    # Extract fields from PDF
    fields = extract_form_fields(pdf_path)
    
    if not fields:
        print("Warning: No form fields found in the PDF.")
    
    # Extract content with field positions
    pages_content = extract_pdf_text_with_fields(pdf_path, fields)
    
    # Create Word document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Form Document', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Track which fields have been added
    added_fields = set()
    
    # Process each page
    for page_num, content_items in enumerate(pages_content):
        if page_num > 0:
            doc.add_page_break()
        
        doc.add_heading(f'Page {page_num + 1}', level=1)
        
        current_paragraph = None
        last_y = -100
        
        for item in content_items:
            if item["type"] == "text":
                # Check if this is a new paragraph (significant vertical gap)
                if abs(item["y"] - last_y) > 20 or current_paragraph is None:
                    current_paragraph = doc.add_paragraph()
                    current_paragraph.add_run(item["content"])
                else:
                    current_paragraph.add_run(" " + item["content"])
                last_y = item["y"]
                
            elif item["type"] == "field":
                field = item["field"]
                field_id = f"{field['name']}_{field['page']}"
                
                if field_id not in added_fields:
                    added_fields.add(field_id)
                    
                    # Add field placeholder
                    p = doc.add_paragraph()
                    
                    # Add field label if available
                    if field["name"]:
                        run = p.add_run(f"{field['name']}: ")
                        run.bold = True
                    
                    # Add placeholder in Copilot Studio format
                    placeholder = field["placeholder"]
                    
                    # Add context based on field type
                    if field["subtype"] == "checkbox":
                        p.add_run(f"{placeholder}")
                        p.add_run(f"  [Checkbox]").italic = True
                    elif field["subtype"] == "radio":
                        p.add_run(f"{placeholder}")
                        p.add_run(f"  [Radio Button]").italic = True
                    elif field["subtype"] == "dropdown" or field["subtype"] == "listbox":
                        p.add_run(f"{placeholder}")
                        if field["options"]:
                            options_str = ", ".join(str(o) for o in field["options"][:5])
                            p.add_run(f"  [Options: {options_str}]").italic = True
                        else:
                            p.add_run(f"  [Dropdown]").italic = True
                    else:
                        p.add_run(f"{placeholder}")
                    
                    last_y = item["y"]
    
    # Add field summary section if requested
    if include_field_list and fields:
        doc.add_page_break()
        doc.add_heading('Form Fields Summary', level=1)
        
        # Create a summary table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Field Name'
        header_cells[1].text = 'Type'
        header_cells[2].text = 'Placeholder'
        header_cells[3].text = 'Page'
        
        # Make header bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Add field rows
        unique_fields = {}
        for field in fields:
            key = field["sanitized_name"]
            if key not in unique_fields:
                unique_fields[key] = field
        
        for field in unique_fields.values():
            row_cells = table.add_row().cells
            row_cells[0].text = field["name"] or "(unnamed)"
            row_cells[1].text = field["subtype"]
            row_cells[2].text = field["placeholder"]
            row_cells[3].text = str(field["page"])
    
    # Save the document
    doc.save(output_path)
    print(f"Word document saved to: {output_path}")
    
    return fields


def main():
    if len(sys.argv) < 2:
        print("Usage: python pdf_form_to_word.py <input.pdf> [output.docx]")
        print("\nConverts PDF form fields to Word document with Copilot Studio placeholders.")
        print("\nPlaceholder format: {{FieldName}}")
        print("For table fields: {{TableName.ColumnName}}")
        sys.exit(1)
    
    input_pdf = sys.argv[1]
    
    if not Path(input_pdf).exists():
        print(f"Error: Input file '{input_pdf}' not found.")
        sys.exit(1)
    
    # Generate output filename if not provided
    if len(sys.argv) >= 3:
        output_docx = sys.argv[2]
    else:
        output_docx = Path(input_pdf).stem + "_form.docx"
    
    print(f"Converting: {input_pdf}")
    print(f"Output: {output_docx}")
    print("-" * 50)
    
    fields = create_word_document(input_pdf, output_docx)
    
    print("-" * 50)
    print(f"Found {len(fields)} form field(s)")
    
    if fields:
        print("\nField placeholders created:")
        seen = set()
        for field in fields:
            if field["placeholder"] not in seen:
                seen.add(field["placeholder"])
                print(f"  - {field['placeholder']} ({field['subtype']})")


if __name__ == "__main__":
    main()
